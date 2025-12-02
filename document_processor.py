#!/usr/bin/env python3

"""
Document Processing Module for RFP Studio

Handles document text extraction, question detection, and answer generation
for the Streamlit RAG frontend.
"""

import re
import os
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import tempfile

# Document processing imports
try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

# RFP Studio imports
from rfp_studio.vector import embed_text, search_knowledge_base
from rfp_studio.config import get_settings

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles text extraction from various document formats."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
        
    def extract_text(self, file_path: str) -> str:
        """Extract text from a document file."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._extract_pdf_text(file_path)
        elif extension == '.docx':
            return self._extract_docx_text(file_path)
        elif extension == '.txt':
            return self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def get_page_count(self, file_path: str) -> int:
        """Get the number of pages in a document."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            try:
                if PdfReader:
                    with open(file_path, 'rb') as file:
                        reader = PdfReader(file)
                        return len(reader.pages)
            except Exception:
                pass
        
        # Default to 1 page for other formats or if PDF reading fails
        return 1
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        if not PdfReader:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")
            
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += f"\\n\\n--- Page {page_num + 1} ---\\n{page_text}"
                    except Exception as e:
                        logger.warning(f"Error extracting text from PDF page {page_num + 1}: {e}")
                        
        except Exception as e:
            raise RuntimeError(f"Error reading PDF file: {e}")
            
        return text.strip()
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        if not Document:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
            
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
                        
            return "\\n\\n".join(text_parts)
            
        except Exception as e:
            raise RuntimeError(f"Error reading DOCX file: {e}")
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise RuntimeError(f"Error reading text file: {e}")


class QuestionExtractor:
    """Extracts questions from document text."""
    
    def __init__(self):
        # Common question patterns
        self.question_patterns = [
            r'\\b\\d+\\.\\s*[^.?!]*\\?',  # Numbered questions (1. What is...?)
            r'\\b[A-Z][^.?!]*\\?\\s*$',  # Sentences ending with ?
            r'(?:Please\\s+)?(?:describe|explain|provide|list|identify|specify|detail)\\s+[^.?!]*[.?]',
            r'(?:What|How|Where|When|Why|Who|Which)\\s+[^.?!]*[.?]',
            r'(?:Do|Does|Can|Could|Will|Would|Should|Is|Are)\\s+[^.?!]*[.?]',
        ]
        
        # Compile patterns
        self.compiled_patterns = [re.compile(pattern, re.IGNORECASE | re.MULTILINE) 
                                 for pattern in self.question_patterns]
        
    def extract_questions(self, text: str) -> List[str]:
        """Extract questions from document text."""
        questions = []
        
        # Clean text first
        clean_text = self._clean_text(text)
        
        # Apply each pattern
        for pattern in self.compiled_patterns:
            matches = pattern.findall(clean_text)
            for match in matches:
                question = self._clean_question(match)
                if question and len(question) > 10:  # Filter out very short questions
                    questions.append(question)
        
        # Remove duplicates while preserving order
        unique_questions = []
        seen = set()
        
        for q in questions:
            q_normalized = q.lower().strip()
            if q_normalized not in seen:
                seen.add(q_normalized)
                unique_questions.append(q)
        
        return unique_questions[:50]  # Limit to first 50 questions
    
    def _clean_text(self, text: str) -> str:
        """Clean and prepare text for question extraction."""
        # Remove excessive whitespace
        text = re.sub(r'\\s+', ' ', text)
        
        # Remove page markers
        text = re.sub(r'--- Page \\d+ ---', '\\n', text)
        
        # Normalize line breaks
        text = re.sub(r'\\n+', '\\n', text)
        
        return text.strip()
    
    def _clean_question(self, question: str) -> str:
        """Clean and format a single question."""
        # Remove leading numbers and bullets
        question = re.sub(r'^\\s*\\d+\\.\\s*', '', question)
        question = re.sub(r'^\\s*[•·-]\\s*', '', question)
        
        # Clean whitespace
        question = re.sub(r'\\s+', ' ', question).strip()
        
        # Ensure question ends with ?
        if not question.endswith('?') and not question.endswith('.'):
            question += '?'
            
        return question


class AnswerGenerator:
    """Generates answers using RAG pipeline."""
    
    def __init__(self):
        self.settings = get_settings()
        
    def generate_answer(self, question: str, confidence_threshold: float = 0.7) -> Dict[str, Any]:
        """Generate answer for a question using RAG."""
        try:
            # Step 1: Embed the question
            question_embedding = embed_text(question)
            
            # Step 2: Search knowledge base
            search_results = search_knowledge_base(question_embedding, limit=3)
            
            if not search_results:
                return {
                    'success': False,
                    'error': 'No relevant knowledge found',
                    'question': question
                }
            
            # Step 3: Check confidence (using search score as proxy)
            top_score = search_results[0].get('score', 0.0) if search_results else 0.0
            
            if top_score < confidence_threshold:
                return {
                    'success': False,
                    'error': f'Confidence too low: {top_score:.3f} < {confidence_threshold}',
                    'question': question,
                    'search_results': search_results
                }
            
            # Step 4: Generate answer from top results
            answer = self._synthesize_answer(question, search_results)
            
            # Step 5: Extract sources
            sources = self._extract_sources(search_results)
            
            return {
                'success': True,
                'question': question,
                'answer': answer,
                'confidence': top_score,
                'sources': sources,
                'search_results': search_results
            }
            
        except Exception as e:
            logger.error(f"Error generating answer: {e}")
            return {
                'success': False,
                'error': str(e),
                'question': question
            }
    
    def _synthesize_answer(self, question: str, search_results: List[Dict[str, Any]]) -> str:
        """Synthesize answer from search results."""
        # Simple approach: use the top result's text
        # In production, you'd use LLM to synthesize multiple sources
        
        if not search_results:
            return "I don't have enough information to answer this question."
        
        top_result = search_results[0]
        text = top_result.get('text', '')
        
        # Try to extract answer-relevant portion
        if 'A:' in text:
            # If it's a Q&A format, extract the answer part
            parts = text.split('A:', 1)
            if len(parts) > 1:
                answer = parts[1].strip()
                if answer:
                    return answer
        
        # Otherwise, use the full text (up to reasonable length)
        max_length = 500
        if len(text) > max_length:
            # Try to end at sentence boundary
            truncated = text[:max_length]
            last_sentence = truncated.rfind('.')
            if last_sentence > max_length * 0.7:  # If we can find a good break point
                return truncated[:last_sentence + 1]
            else:
                return truncated + "..."
        
        return text
    
    def _extract_sources(self, search_results: List[Dict[str, Any]]) -> List[str]:
        """Extract source information from search results."""
        sources = []
        
        for result in search_results:
            team = result.get('team_key', 'Unknown Team')
            topic = result.get('topic', 'General Knowledge')
            score = result.get('score', 0.0)
            
            source = f"{team} - {topic} (score: {score:.3f})"
            sources.append(source)
        
        return sources


def validate_dependencies() -> Dict[str, bool]:
    """Check which document processing dependencies are available."""
    deps = {
        'PyPDF2': PyPDF2 is not None,
        'python-docx': Document is not None,
        'PIL': Image is not None,
        'pytesseract': pytesseract is not None
    }
    
    return deps


def get_processing_capabilities() -> Dict[str, Any]:
    """Get information about processing capabilities."""
    deps = validate_dependencies()
    
    supported_formats = ['txt']  # Always supported
    
    if deps['PyPDF2']:
        supported_formats.append('pdf')
        
    if deps['python-docx']:
        supported_formats.append('docx')
    
    return {
        'supported_formats': supported_formats,
        'dependencies': deps,
        'ocr_available': deps['PIL'] and deps['pytesseract']
    }


if __name__ == "__main__":
    # Test the processing capabilities
    print("Document Processing Capabilities:")
    caps = get_processing_capabilities()
    
    print(f"Supported formats: {', '.join(caps['supported_formats'])}")
    print(f"Dependencies: {caps['dependencies']}")
    print(f"OCR available: {caps['ocr_available']}")